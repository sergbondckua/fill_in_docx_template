import locale
import platform
from datetime import datetime
from dataclasses import dataclass
from docx import Document
from declension import NameDeclension
from numwords import FinancialAmountInUAH

# Встановлюємо українську локаль
if platform.system() == "Windows":
    locale.setlocale(locale.LC_TIME, "Ukrainian")
else:
    locale.setlocale(locale.LC_TIME, "uk_UA.UTF-8")


@dataclass
class PartyData:
    contract_number: str
    full_name: str
    short_name: str
    address: str
    person_name: str
    phone_number: str
    city: str
    bank_details: str


@dataclass
class ContractGenerator:
    source_price: float
    from_date: str = datetime.now().replace(day=1).strftime('"%d" %B %Y')

    def __post_init__(self):
        self.price = int(self.source_price)
        self.pennies = FinancialAmountInUAH(
            self.source_price
        ).extract_pennies()
        self.total_price = self.source_price * 12
        self.total_pennies = FinancialAmountInUAH(
            self.total_price
        ).extract_pennies()
        self.price_text = FinancialAmountInUAH(
            self.source_price
        ).format_result()
        self.total_price_text = FinancialAmountInUAH(
            self.total_price
        ).format_result()
        self.genitive_name = NameDeclension()

    def get_contract_data(self, party_data: PartyData):
        part_person_name = party_data.person_name.split()
        short_name = f"{part_person_name[1]} {part_person_name[0].upper()}"
        return {
            "contract_number": party_data.contract_number,
            "city": party_data.city,
            "from_date": self.from_date,
            "party_one": party_data.full_name.upper(),
            "party_one_short_name": (
                party_data.short_name.upper()
                if party_data.short_name
                else party_data.full_name.upper()
            ),
            "person_party_one": party_data.person_name,
            "short_name": short_name,
            "genitive_name": self.genitive_name.to_genitive(
                party_data.person_name
            ),
            "address": party_data.address,
            "price": str(self.price),
            "pennies": f"{self.pennies}0"[:2],
            "price_text": self.price_text,
            "total_price_text": self.total_price_text,
            "total_price": f"{int(self.total_price)}",
            "total_pennies": f"{self.total_pennies}0"[:2],
            "person_party_one_phonenumber": party_data.phone_number,
            "bank_details": party_data.bank_details,
        }


class TemplateFiller:
    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path

    def fill_template(self, data, bold_keys=None):
        if bold_keys is None:
            bold_keys = []

        doc = Document(self.template_path)

        def replace_markers(text, marker_key, marker_value, is_bold):
            if f"{{{marker_key}}}" in text:
                text = text.replace(f"{{{marker_key}}}", marker_value)
                if is_bold:
                    return text, True
            return text, False

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run_text = run.text
                for key, value in data.items():
                    run_text, should_bold = replace_markers(
                        run_text, key, value, key in bold_keys
                    )
                    run.text = run_text
                    if should_bold:
                        run.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            for key, value in data.items():
                                run_text, should_bold = replace_markers(
                                    run_text, key, value, key in bold_keys
                                )
                                run.text = run_text
                                if should_bold:
                                    run.bold = True

        doc.save(self.output_path)


if __name__ == "__main__":

    # Приклад використання
    party_data_src = PartyData(
        contract_number=f"ОСББ-{datetime.now().replace(day=1).strftime('%d%m%Y')}-105",
        full_name='ОБ\'ЄДНАННЯ СПІВВЛАСНИКІВ БАГАТОКВАРТИРНОГО БУДИНКУ ЮГОСЛАВСЬКИЙ',
        short_name='ОСББ "ЮГОСЛАВСЬКИЙ"',
        address="вулиця Пастерівська, будинок 11",
        person_name="Пінчук Микола Васильович",
        phone_number="096-404-98-06",
        city="Черкаси",
        bank_details="""18016, м.Черкаси, вул.Пастерівська , 11
р/р UA343052990000026002021604334
в АТ КБ "Приватбанк"
МФО 305299
ЄДРПОУ 40726671
т. 097-507-40-46
""",
    )

    contract = ContractGenerator(source_price=7100*0.05)
    contract_data = contract.get_contract_data(party_data_src)

    template_contract_filler = TemplateFiller(
        "templates/contract_template.docx", "filled_contract.docx"
    )
    template_pax_akt_filler = TemplateFiller(
        "templates/pax_akt_template.docx", "filled_rax_akt.docx"
    )
    template_contract_filler.fill_template(contract_data)
    template_pax_akt_filler.fill_template(contract_data)

from docx import Document


class TemplateFiller:
    """Заповнює шаблон договору"""

    def __init__(self, template_path, output_path):
        self.template_path = template_path
        self.output_path = output_path

    def fill_template(self, data: dict, bold_keys: list = None):
        """Заповнення шаблону договору"""

        # Перевіряємо, чи передано маркери для жирного тексту
        if bold_keys is None:
            bold_keys = []

        # Шаблон договору
        doc = Document(self.template_path)

        def replace_markers(
            text: str, marker_key: str, marker_value: str, is_bold: bool
        ):
            """Замінює маркери в тексті на відповідні значення"""

            # Перевіряє, чи в тексті є маркер
            if f"{{{marker_key}}}" in text:
                text = text.replace(f"{{{marker_key}}}", marker_value)
                if is_bold:
                    return text, True
            return text, False

        # Замінює маркери в тексті
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run_text = run.text
                for key, value in data.items():
                    run_text, should_bold = replace_markers(
                        run_text, key, value, key in bold_keys
                    )
                    run.text = run_text
                    if should_bold:
                        run.bold = True

        # Замінює маркери в таблицях
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run_text = run.text
                            for key, value in data.items():
                                run_text, should_bold = replace_markers(
                                    run_text, key, value, key in bold_keys
                                )
                                run.text = run_text
                                if should_bold:
                                    run.bold = True

        # Зберігає договір у файл
        doc.save(self.output_path)
